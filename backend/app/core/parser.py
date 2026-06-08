"""合同文件解析器 — 支持 PDF / DOCX / TXT"""

import io
import logging
from pathlib import Path
from typing import NamedTuple

logger = logging.getLogger(__name__)


class ParseResult(NamedTuple):
    text: str
    page_count: int | None  # PDF 有页码，其他为 None


class ContractParser:
    """多格式合同文件解析器"""

    SUPPORTED_TYPES = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "txt": "text/plain",
    }

    @classmethod
    def parse(cls, file_bytes: bytes, file_type: str) -> ParseResult:
        file_type = file_type.lower()
        t0 = __import__("time").perf_counter()
        if file_type == "pdf":
            result = cls._parse_pdf(file_bytes)
        elif file_type in ("docx", "doc"):
            result = cls._parse_docx(file_bytes)
        elif file_type == "txt":
            result = cls._parse_txt(file_bytes)
        else:
            raise ValueError(
                f"不支持的文件格式: {file_type}。支持: {list(cls.SUPPORTED_TYPES.keys())}"
            )
        elapsed = __import__("time").perf_counter() - t0
        logger.info(
            "文件解析完成 | type=%s | size=%dKB | text_chars=%d | %.2fs",
            file_type, len(file_bytes) // 1024, len(result.text), elapsed,
        )
        return result

    @classmethod
    def _parse_pdf(cls, file_bytes: bytes) -> ParseResult:
        import pdfplumber

        text_parts = []
        page_count = 0
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        if not text_parts:
            raise ValueError("无法从 PDF 中提取文本，文件可能是扫描件或图片。")

        return ParseResult(text="\n\n".join(text_parts), page_count=page_count)

    @classmethod
    def _parse_docx(cls, file_bytes: bytes) -> ParseResult:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    text_parts.append(row_text)

        return ParseResult(text="\n".join(text_parts), page_count=None)

    @classmethod
    def _parse_txt(cls, file_bytes: bytes) -> ParseResult:
        # 尝试多种编码
        text = None
        for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            raise ValueError("无法解码 TXT 文件，请确认文件编码为 UTF-8 或 GBK。")

        return ParseResult(text=text, page_count=None)
