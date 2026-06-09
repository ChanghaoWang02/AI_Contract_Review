"""合同 DOCX 导出渲染器"""

import io
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

_FONT_NAME = "Microsoft YaHei"
_FONT_SIZE = Pt(11)
_TITLE_SIZE = Pt(16)
_MARGIN = Cm(2.5)


def _set_cn_font(run, font_name: str = _FONT_NAME):
    """为 run 设置中文字体"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)


def _setup_default_style(doc: Document):
    """设置默认段落样式支持中文"""
    style = doc.styles['Normal']
    style.font.name = _FONT_NAME
    style.font.size = _FONT_SIZE
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # 在 Normal 样式的 rPr 中设置中文字体
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = style.element.makeelement(qn('w:rPr'), {})
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), _FONT_NAME)


def render_contract_docx(content: str, title: str) -> bytes:
    """将合同纯文本渲染为 DOCX 字节"""
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = _MARGIN
        section.bottom_margin = _MARGIN
        section.left_margin = _MARGIN
        section.right_margin = _MARGIN

    _setup_default_style(doc)

    # ── 标题 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = _TITLE_SIZE
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    _set_cn_font(title_run)

    # 标题后空一行
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    # ── 正文 ──
    # 按「第X条」分割，条款标题加粗
    CLAUSE_HEAD_RE = re.compile(r'^第[一二三四五六七八九十百千]+条\b')

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if CLAUSE_HEAD_RE.match(line):
            # 条款标题行：加粗
            run = para.add_run(line)
            run.font.bold = True
            run.font.size = _FONT_SIZE
            _set_cn_font(run)
        else:
            run = para.add_run(line)
            run.font.size = _FONT_SIZE
            _set_cn_font(run)

    # ── 保存 ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
