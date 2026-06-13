"""合同翻译 API — SSE 流式翻译 + 文本翻译 + 保存 + 导出"""

import io
import json
import re
import logging
from pydantic import BaseModel
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse, Response
from sqlalchemy.orm import Session

from app.db.session import get_db
from app.models import Contract
from app.schemas.translate import (
    TranslateGenerateRequest,
    TranslateClauseRequest,
    TranslateTextRequest,
    TranslateSaveRequest,
)
from app.core.chunker import ContractChunker
from app.core.translator import TranslationEngine

router = APIRouter()
logger = logging.getLogger(__name__)

# SSE 数据中不允许的控制字符（U+0000-U+001F 除 \t\n\r）
_SSE_CTRL_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')


def _sanitize_str(s: str) -> str:
    return _SSE_CTRL_RE.sub('', s)


def _sanitize(obj):
    """递归清理 dict/list/str 中的控制字符"""
    if isinstance(obj, str):
        return _sanitize_str(obj)
    if isinstance(obj, dict):
        return {k: _sanitize(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_sanitize(v) for v in obj]
    return obj


def _sse_encode(data: dict) -> str:
    """安全 SSE 编码：清理控制字符并序列化为 JSON"""
    return json.dumps(_sanitize(data), ensure_ascii=False)


@router.post("/generate")
async def translate_generate(body: TranslateGenerateRequest, db: Session = Depends(get_db)):
    """SSE 流式翻译合同全文（逐条条款）"""
    contract = db.query(Contract).filter(Contract.id == body.contract_id).first()
    if not contract:
        raise HTTPException(404, "合同不存在。")
    if not contract.content:
        raise HTTPException(400, "合同内容为空，无法翻译。")

    target_lang = body.target_lang or "zh"

    async def event_generator():
        try:
            # 1. 语言检测
            detect_result = await TranslationEngine.detect_language(
                contract.content, body.provider
            )
            source_lang = detect_result["detected"]
            tier = detect_result["tier"]
            yield f"data: {_sse_encode({'event': 'progress', 'data': {'stage': 'detect', 'detected': source_lang, 'tier': tier, 'language_name': detect_result['language_name']}})}\n\n"

            # 2. 条款分块
            clauses = ContractChunker.split(contract.content)
            if not clauses:
                yield f"data: {_sse_encode({'event': 'error', 'data': '无法解析合同条款'})}\n\n"
                return
            yield f"data: {_sse_encode({'event': 'progress', 'data': {'stage': 'chunked', 'total_clauses': len(clauses)}})}\n\n"

            # 3. 逐条流式翻译
            skipped_clauses = []
            translated_count = 0

            for clause in clauses:
                try:
                    async for event in TranslationEngine.translate_clause_stream(
                        clause=clause,
                        source_lang=source_lang,
                        target_lang=target_lang,
                        tier=tier,
                        total_clauses=len(clauses),
                        provider=body.provider,
                        model=body.model or "",
                    ):
                        if event["event"] == "token":
                            yield f"data: {_sse_encode({'event': 'token', 'data': {'clause_index': clause.index, 'content': event['data']}})}\n\n"
                        elif event["event"] == "clause_done":
                            translated_count += 1
                            yield f"data: {_sse_encode({'event': 'clause_done', 'data': event['data']})}\n\n"
                        elif event["event"] == "clause_error":
                            skipped_clauses.append(event["data"])
                            yield f"data: {_sse_encode({'event': 'clause_error', 'data': event['data']})}\n\n"
                except Exception as e:
                    skipped_clauses.append({
                        "clause_index": clause.index,
                        "clause_id": clause.id,
                        "error": str(e),
                    })
                    yield f"data: {_sse_encode({'event': 'clause_error', 'data': {'clause_index': clause.index, 'clause_id': clause.id, 'error': str(e)}})}\n\n"

            # 4. 完成
            yield f"data: {_sse_encode({'event': 'done', 'data': {'total_clauses': len(clauses), 'translated': translated_count, 'skipped_clauses': skipped_clauses, 'source_lang': source_lang, 'target_lang': target_lang}})}\n\n"

        except Exception as e:
            logger.error("合同翻译全局异常: %s", e, exc_info=True)
            yield f"data: {_sse_encode({'event': 'error', 'data': str(e)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream; charset=utf-8")


@router.post("/clause")
async def translate_clause(body: TranslateClauseRequest, db: Session = Depends(get_db)):
    """SSE 流式重新翻译单条条款"""
    contract = db.query(Contract).filter(Contract.id == body.contract_id).first()
    if not contract:
        raise HTTPException(404, "合同不存在。")

    # 检测语言
    detect_result = await TranslationEngine.detect_language(body.original_text)
    source_lang = detect_result["detected"]
    tier = detect_result["tier"]

    # 目标语言：来自请求参数（前端选择）
    target_lang = body.target_lang

    # 构造 Clause 对象
    clause = ContractChunker._build_clause(
        title=f"条款 {body.clause_index + 1}",
        content=body.original_text,
        index=body.clause_index,
    )

    async def event_generator():
        try:
            async for event in TranslationEngine.translate_clause_stream(
                clause=clause,
                source_lang=source_lang,
                target_lang=target_lang,
                tier=tier,
                total_clauses=1,
                provider=None,
                model="",
                instruction=body.instruction,
            ):
                if event["event"] == "token":
                    yield f"data: {_sse_encode({'event': 'token', 'data': event['data']})}\n\n"
                elif event["event"] == "clause_done":
                    yield f"data: {_sse_encode({'event': 'done', 'data': event['data']})}\n\n"
                elif event["event"] == "clause_error":
                    yield f"data: {_sse_encode({'event': 'error', 'data': event['data']})}\n\n"
        except Exception as e:
            yield f"data: {_sse_encode({'event': 'error', 'data': str(e)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream; charset=utf-8")


@router.post("/text")
async def translate_text(body: TranslateTextRequest):
    """SSE 流式翻译任意文本（审核报告等）"""
    async def event_generator():
        try:
            async for event in TranslationEngine.translate_text_stream(
                content=body.content,
                target_lang=body.target_lang,
                source_lang=body.source_lang,
                provider=body.provider,
                model=body.model or "",
            ):
                if event["event"] == "progress":
                    yield f"data: {_sse_encode({'event': 'progress', 'data': event['data']})}\n\n"
                elif event["event"] == "token":
                    yield f"data: {_sse_encode({'event': 'token', 'data': event['data']})}\n\n"
                elif event["event"] == "done":
                    yield f"data: {_sse_encode({'event': 'done', 'data': event['data']})}\n\n"
                elif event["event"] == "error":
                    yield f"data: {_sse_encode({'event': 'error', 'data': event['data']})}\n\n"
        except Exception as e:
            logger.error("文本翻译全局异常: %s", e, exc_info=True)
            yield f"data: {_sse_encode({'event': 'error', 'data': str(e)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream; charset=utf-8")


@router.post("/save")
async def translate_save(body: TranslateSaveRequest, db: Session = Depends(get_db)):
    """保存译文为新合同（子合同），可直接审核"""
    parent = db.query(Contract).filter(Contract.id == body.contract_id).first()
    if not parent:
        raise HTTPException(404, "原文合同不存在。")

    # 构建文件名
    clean_name = body.filename
    for ext in (".txt", ".pdf", ".docx"):
        if clean_name.lower().endswith(ext):
            clean_name = clean_name.rsplit(".", 1)[0]
            break
    if not clean_name.endswith(f"_{body.target_lang.upper()}"):
        clean_name = f"{clean_name}_{body.target_lang.upper()}"

    # 分块统计
    clause_count = len(ContractChunker.split(body.translated_content))

    child = Contract(
        filename=f"{clean_name}.txt",
        original_filename=f"{clean_name}.txt",
        content=body.translated_content,
        content_type="txt",
        source="translated",
        parent_contract_id=body.contract_id,
        source_lang=body.source_lang,
        target_lang=body.target_lang,
        file_size=len(body.translated_content.encode("utf-8")),
        clause_count=clause_count,
    )
    db.add(child)
    db.commit()
    db.refresh(child)

    logger.info(
        "译文已保存: id=%d parent=%d lang=%s→%s clauses=%d",
        child.id, body.contract_id, body.source_lang, body.target_lang, clause_count,
    )

    return {
        "id": child.id,
        "parent_contract_id": body.contract_id,
        "source": "translated",
        "source_lang": body.source_lang,
        "target_lang": body.target_lang,
        "clause_count": clause_count,
        "filename": child.original_filename,
    }


def _build_translate_pdf(content: str, filename: str) -> bytes:
    """将译文生成为 PDF 文件（字节）"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from app.core.pdf_renderer import FONT_NAME

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=25*2.83, rightMargin=25*2.83,
                            topMargin=20*2.83, bottomMargin=20*2.83)

    styles = {
        "title": ParagraphStyle("TTitle", fontName=FONT_NAME, fontSize=18,
                                leading=26, alignment=TA_CENTER, spaceAfter=20),
        "body": ParagraphStyle("TBody", fontName=FONT_NAME, fontSize=10,
                               leading=18, spaceAfter=10),
    }

    story = []
    # 标题
    clean = filename
    for ext in (".txt", ".pdf", ".docx"):
        if clean.lower().endswith(ext):
            clean = clean.rsplit(".", 1)[0]
            break
    story.append(Paragraph(f"合同译文：{clean}", styles["title"]))
    story.append(Spacer(1, 12))

    # 正文：按段落拆分
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    for para in paragraphs:
        # 将换行转为 <br/>
        text = para.replace('\n', '<br/>')
        story.append(Paragraph(text, styles["body"]))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def _build_translate_docx(content: str) -> bytes:
    """将译文生成为 DOCX 文件（字节）"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 正文
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    for i, para in enumerate(paragraphs):
        p = doc.add_paragraph()
        p.style.font.size = Pt(11)
        p.style.font.name = 'Microsoft YaHei'
        # 拆分内部换行
        lines = para.split('\n')
        for j, line in enumerate(lines):
            if j > 0:
                p.add_run('\n').font.size = Pt(11)
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.name = 'Microsoft YaHei'

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class TranslateExportRequest(BaseModel):
    """译文导出请求"""
    content: str
    format: str = "pdf"   # pdf / docx / txt
    filename: str = "译文"


@router.post("/export")
async def export_translation(body: TranslateExportRequest):
    """导出译文为指定格式（PDF / DOCX / TXT）

    接受前端传来的译文内容，直接转换格式后返回文件。
    无需依赖数据库——用户可能在保存前就导出。
    """
    if not body.content:
        raise HTTPException(400, "译文内容为空，无法导出。")

    if body.format not in ("pdf", "docx", "txt"):
        raise HTTPException(400, f"不支持的导出格式：{body.format}，支持 pdf/docx/txt")

    clean_name = body.filename
    for ext in (".txt", ".pdf", ".docx"):
        if clean_name.lower().endswith(ext):
            clean_name = clean_name.rsplit(".", 1)[0]
            break

    if body.format == "txt":
        content_bytes = body.content.encode("utf-8")
        media_type = "text/plain; charset=utf-8"
        return Response(
            content=content_bytes,
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{clean_name}.txt"'},
        )

    if body.format == "pdf":
        try:
            pdf_bytes = _build_translate_pdf(body.content, clean_name)
        except Exception as e:
            logger.error("PDF 生成失败: %s", e, exc_info=True)
            raise HTTPException(500, f"PDF 生成失败：{e}")
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{clean_name}.pdf"'},
        )

    if body.format == "docx":
        try:
            docx_bytes = _build_translate_docx(body.content)
        except Exception as e:
            logger.error("DOCX 生成失败: %s", e, exc_info=True)
            raise HTTPException(500, f"DOCX 生成失败：{e}")
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{clean_name}.docx"'},
        )
